from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REVIEW_DIR = ROOT / "docs" / "templates" / "document-defaults"
SEED_DIR = ROOT / "NexusFlow.Web" / "SeedData" / "DocumentTemplates" / "Defaults"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(94, 108, 132)
WHITE = RGBColor(255, 255, 255)
LIGHT_FILL = "F2F4F7"
HEADER_FILL = "E8EEF5"
BORDER = "CBD5E1"


STANDARD_LINE_COLUMNS = [
    ("Code", "ItemCode", 950, "left"),
    ("Description", "Description", 2900, "left"),
    ("Qty", "Quantity", 800, "right"),
    ("Unit", "Unit", 650, "left"),
    ("Unit Price", "UnitPrice", 1100, "right"),
    ("Discount", "Discount", 900, "right"),
    ("Tax", "TaxAmount", 850, "right"),
    ("Line Total", "LineTotal", 1210, "right"),
]


DOCUMENTS = [
    {
        "type": "SalesOrder",
        "title": "SALES ORDER",
        "party": "Customer",
        "billing": "Billing Address",
        "shipping": "Delivery Address",
        "terms": "Sales order terms: availability, delivery dates, and prices are subject to final confirmation unless otherwise agreed in writing.",
        "payment": False,
        "signatures": ("Prepared By", "Customer Acceptance"),
        "table_title": "Sales Lines",
        "table_name": "SalesLines",
        "columns": STANDARD_LINE_COLUMNS,
    },
    {
        "type": "SalesQuotation",
        "title": "SALES QUOTATION",
        "party": "Customer",
        "billing": "Billing Address",
        "shipping": "Delivery Address",
        "terms": "Quotation terms: this quotation is editable by the administrator and should be updated with validity, delivery, and warranty conditions.",
        "payment": False,
        "signatures": ("Prepared By", "Approved By"),
        "table_title": "Quotation Lines",
        "table_name": "SalesLines",
        "columns": STANDARD_LINE_COLUMNS,
    },
    {
        "type": "SalesInvoice",
        "title": "SALES INVOICE",
        "party": "Customer",
        "billing": "Billing Address",
        "shipping": "Shipping Address",
        "terms": "Invoice terms: payment is due according to the agreed commercial terms. Update this placeholder with tenant-specific terms.",
        "payment": True,
        "signatures": ("Prepared By", "Received By"),
        "table_title": "Invoice Lines",
        "table_name": "SalesLines",
        "columns": STANDARD_LINE_COLUMNS,
    },
    {
        "type": "CreditNote",
        "title": "CREDIT NOTE",
        "party": "Customer",
        "billing": "Customer Address",
        "shipping": "Reference Address",
        "terms": "Credit note terms: this document adjusts the referenced customer balance subject to approval and posting controls.",
        "payment": False,
        "signatures": ("Prepared By", "Approved By"),
        "table_title": "Returned / Credited Lines",
        "table_name": "CreditNoteLines",
        "columns": [
            ("Code", "ItemCode", 1050, "left"),
            ("Description", "Description", 3500, "left"),
            ("Returned Qty", "Quantity", 1200, "right"),
            ("Unit Price", "UnitPrice", 1200, "right"),
            ("Tax", "TaxAmount", 900, "right"),
            ("Credit Total", "LineTotal", 1510, "right"),
        ],
    },
    {
        "type": "PurchaseOrder",
        "title": "PURCHASE ORDER",
        "party": "Supplier",
        "billing": "Supplier Address",
        "shipping": "Deliver To",
        "terms": "Purchase order terms: supplier must reference this PO number on all invoices, delivery notes, and correspondence.",
        "payment": False,
        "signatures": ("Prepared By", "Supplier Acceptance"),
        "table_title": "Purchase Lines",
        "table_name": "PurchaseLines",
        "columns": [
            ("Code", "ItemCode", 1050, "left"),
            ("Description", "Description", 3300, "left"),
            ("Ordered Qty", "Quantity", 1100, "right"),
            ("Unit", "Unit", 750, "left"),
            ("Unit Cost", "UnitPrice", 1300, "right"),
            ("Line Total", "LineTotal", 1860, "right"),
        ],
    },
    {
        "type": "GRN",
        "title": "GOODS RECEIVED NOTE",
        "party": "Supplier",
        "billing": "Supplier Address",
        "shipping": "Receiving Location",
        "terms": "GRN terms: receipt is subject to quantity and quality inspection. Variances must be recorded through the ERP workflow.",
        "payment": False,
        "signatures": ("Received By", "Checked By"),
        "table_title": "Received Lines",
        "table_name": "ReceivedLines",
        "columns": [
            ("Code", "ItemCode", 1050, "left"),
            ("Description", "Description", 3400, "left"),
            ("Received Qty", "Quantity", 1250, "right"),
            ("Unit", "Unit", 700, "left"),
            ("Unit Cost", "UnitPrice", 1300, "right"),
            ("Value", "LineTotal", 1660, "right"),
        ],
    },
    {
        "type": "SupplierBill",
        "title": "SUPPLIER BILL",
        "party": "Supplier",
        "billing": "Supplier Address",
        "shipping": "Bill Reference",
        "terms": "Supplier bill terms: amounts are payable only after matching, approval, and posting according to company policy.",
        "payment": True,
        "signatures": ("Prepared By", "Approved By"),
        "table_title": "Bill Lines",
        "table_name": "SupplierBillLines",
        "columns": STANDARD_LINE_COLUMNS,
    },
    {
        "type": "DebitNote",
        "title": "DEBIT NOTE",
        "party": "Supplier",
        "billing": "Supplier Address",
        "shipping": "Reference Address",
        "terms": "Debit note terms: this document adjusts the supplier balance subject to approval and posting controls.",
        "payment": False,
        "signatures": ("Prepared By", "Approved By"),
        "table_title": "Debit Note Lines",
        "table_name": "DebitNoteLines",
        "columns": [
            ("Code", "ItemCode", 1050, "left"),
            ("Description", "Description", 3600, "left"),
            ("Qty", "Quantity", 800, "right"),
            ("Unit Price", "UnitPrice", 1200, "right"),
            ("Tax", "TaxAmount", 900, "right"),
            ("Debit Total", "LineTotal", 1810, "right"),
        ],
    },
    {
        "type": "CustomerReceipt",
        "title": "CUSTOMER RECEIPT",
        "party": "Customer",
        "billing": "Customer Address",
        "shipping": "Receipt Reference",
        "terms": "Receipt terms: this receipt confirms payment received and should be matched to the relevant invoices in the ERP.",
        "payment": True,
        "signatures": ("Received By", "Customer Signature"),
        "table_title": "Invoice Allocations",
        "table_name": "PaymentAllocations",
        "columns": [
            ("Reference No", "ReferenceNumber", 2200, "left"),
            ("Allocation Description", "Description", 4300, "left"),
            ("Allocated Amount", "Amount", 2860, "right"),
        ],
    },
    {
        "type": "SupplierPaymentRemittance",
        "title": "SUPPLIER PAYMENT REMITTANCE",
        "party": "Supplier",
        "billing": "Supplier Address",
        "shipping": "Payment Reference",
        "terms": "Remittance terms: payment allocation details are shown for supplier reconciliation. Update banking text as required.",
        "payment": True,
        "signatures": ("Prepared By", "Authorized By"),
        "table_title": "Bill Allocations",
        "table_name": "PaymentAllocations",
        "columns": [
            ("Reference No", "ReferenceNumber", 2200, "left"),
            ("Allocation Description", "Description", 4300, "left"),
            ("Paid Amount", "Amount", 2860, "right"),
        ],
    },
    {
        "type": "StockTransferDeliveryNote",
        "title": "STOCK TRANSFER DELIVERY NOTE",
        "party": "Transfer",
        "billing": "Source Location",
        "shipping": "Destination Location",
        "terms": "Transfer terms: stock movement must be checked at dispatch and receiving locations before final acceptance.",
        "payment": False,
        "signatures": ("Dispatched By", "Received By"),
        "table_title": "Transfer Lines",
        "table_name": "TransferLines",
        "columns": [
            ("Code", "ItemCode", 1100, "left"),
            ("Description", "Description", 4000, "left"),
            ("Transfer Qty", "Quantity", 1300, "right"),
            ("Unit", "Unit", 800, "left"),
            ("Unit Cost", "UnitPrice", 1100, "right"),
            ("Value", "LineTotal", 1060, "right"),
        ],
    },
    {
        "type": "Payslip",
        "title": "PAYSLIP",
        "party": "Employee",
        "billing": "Employee Details",
        "shipping": "Pay Period / Work Location",
        "terms": "Payslip note: earnings, deductions, and net pay are generated from payroll records. Update statutory text as required.",
        "payment": False,
        "signatures": ("Prepared By", "Employee Acknowledgement"),
        "payroll": True,
    },
]


def set_run_font(run, size=9, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_twips=9360):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def clear_cell(cell):
    cell.text = ""
    return cell.paragraphs[0]


def add_merge_field(paragraph, field_name, display=None, size=9, color=None, bold=False, hidden=False):
    if display is None:
        display = f"{{{{{field_name}}}}}"

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" MERGEFIELD  {field_name}  \\* MERGEFORMAT "
    run._r.append(instr)

    run = paragraph.add_run()
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    result = paragraph.add_run(display)
    set_run_font(result, size=size, color=color, bold=bold)
    result.font.hidden = hidden

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    return result


def add_static_text(paragraph, text, size=9, color=None, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def style_paragraph(paragraph, before=0, after=2, line_spacing=1.05, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if align is not None:
        paragraph.alignment = align


def add_label_value(cell, label, field_name, display=None):
    p = clear_cell(cell)
    p.paragraph_format.space_after = Pt(1)
    add_static_text(p, label + ": ", size=8.2, color=MUTED, bold=True)
    add_merge_field(p, field_name, display=display, size=8.2, color=NAVY, bold=True)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(13)
    section.bottom_margin = Mm(13)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05


def add_header(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table)
    for index, width in enumerate((2700, 6660)):
        cell = table.cell(0, index)
        set_cell_width(cell, width)
        set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    left = clear_cell(table.cell(0, 0))
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_merge_field(left, "Image:CompanyLogo", display="[Company Logo]", size=8, color=MUTED)

    right = clear_cell(table.cell(0, 1))
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_merge_field(right, "CompanyName", size=14, color=NAVY, bold=True)

    for field, label in (
        ("CompanyAddress", None),
        ("CompanyEmail", "Email"),
        ("CompanyPhone", "Phone"),
        ("CompanyTaxRegistrationNumber", "Tax Reg"),
        ("CompanyBusinessRegistrationNumber", "Business Reg"),
    ):
        p = table.cell(0, 1).add_paragraph()
        style_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
        if label:
            add_static_text(p, label + ": ", size=7.5, color=MUTED, bold=True)
        add_merge_field(p, field, size=7.5, color=MUTED)

    set_table_borders(table, color="FFFFFF", size="0")

    p = doc.add_paragraph()
    style_paragraph(p, before=2, after=4)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def add_title_block(doc, config):
    table = doc.add_table(rows=2, cols=4)
    set_table_width(table)
    set_table_borders(table)
    for row in table.rows:
        for index, width in enumerate((3500, 2200, 1600, 2060)):
            set_cell_width(row.cells[index], width)
            set_cell_margins(row.cells[index], top=80, bottom=80, start=120, end=120)

    title_cell = table.cell(0, 0)
    title_cell.merge(table.cell(1, 0))
    set_cell_fill(title_cell, "0B2545")
    p = clear_cell(title_cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_static_text(p, config["title"], size=16, color=WHITE, bold=True)

    add_label_value(table.cell(0, 1), "Document No", "DocumentNumber")
    add_label_value(table.cell(0, 2), "Date", "DocumentDate")
    add_label_value(table.cell(0, 3), "Currency", "CurrencyCode")
    add_label_value(table.cell(1, 1), config["party"], "CustomerOrSupplierName")
    ref_cell = table.cell(1, 2)
    ref_cell.merge(table.cell(1, 3))
    p = clear_cell(ref_cell)
    add_static_text(p, "Template Type: ", size=8.2, color=MUTED, bold=True)
    add_static_text(p, config["type"], size=8.2, color=NAVY, bold=True)


def add_address_block(doc, config):
    table = doc.add_table(rows=2, cols=2)
    set_table_width(table)
    set_table_borders(table)
    for row in table.rows:
        for index, width in enumerate((4680, 4680)):
            set_cell_width(row.cells[index], width)
            set_cell_margins(row.cells[index], top=80, bottom=80, start=120, end=120)

    for index, label in enumerate((config["billing"], config["shipping"])):
        cell = table.cell(0, index)
        set_cell_fill(cell, LIGHT_FILL)
        p = clear_cell(cell)
        add_static_text(p, label.upper(), size=8, color=NAVY, bold=True)

    for index, field in enumerate(("BillingAddress", "ShippingAddress")):
        p = clear_cell(table.cell(1, index))
        add_merge_field(p, field, size=8.5)


def add_repeating_table(doc, title, group_name, columns):
    p = doc.add_paragraph()
    style_paragraph(p, before=4, after=2)
    add_static_text(p, title, size=10.5, color=NAVY, bold=True)

    table = doc.add_table(rows=2, cols=len(columns))
    set_table_width(table)
    set_table_borders(table)

    for index, (header, field, width, align) in enumerate(columns):
        cell = table.cell(0, index)
        set_cell_width(cell, width)
        set_cell_fill(cell, HEADER_FILL)
        set_cell_margins(cell, top=70, bottom=70, start=80, end=80)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_static_text(p, header, size=7.5, color=NAVY, bold=True)

        cell = table.cell(1, index)
        set_cell_width(cell, width)
        set_cell_margins(cell, top=70, bottom=70, start=80, end=80)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == "right" else WD_ALIGN_PARAGRAPH.LEFT
        if index == 0:
            add_merge_field(p, f"TableStart:{group_name}", display="", size=1, hidden=True)
        add_merge_field(p, field, size=7.7)
        if index == len(columns) - 1:
            add_merge_field(p, f"TableEnd:{group_name}", display="", size=1, hidden=True)


def add_payslip_tables(doc):
    p = doc.add_paragraph()
    style_paragraph(p, before=4, after=2)
    add_static_text(p, "Payroll Summary", size=10.5, color=NAVY, bold=True)

    summary = doc.add_table(rows=4, cols=4)
    set_table_width(summary)
    set_table_borders(summary)
    rows = [
        ("Employee Code", "EmployeeCode", "Pay Period", "PayPeriod"),
        ("Gross Basic", "GrossBasic", "Allowances", "TotalAllowances"),
        ("Total Earnings", "TotalEarnings", "Deductions", "TotalDeductions"),
        ("Net Pay", "NetPay", "Currency", "CurrencyCode"),
    ]
    for row_index, row_values in enumerate(rows):
        for col_index, value in enumerate(row_values):
            cell = summary.cell(row_index, col_index)
            set_cell_width(cell, 2340)
            set_cell_margins(cell, top=70, bottom=70, start=100, end=100)
            if col_index % 2 == 0:
                set_cell_fill(cell, LIGHT_FILL)
                p = clear_cell(cell)
                add_static_text(p, value, size=8, color=MUTED, bold=True)
            else:
                p = clear_cell(cell)
                add_merge_field(p, value, size=8, color=NAVY, bold=row_index == 3 and col_index == 1)

    wrapper = doc.add_table(rows=1, cols=2)
    set_table_width(wrapper)
    set_table_borders(wrapper, color="FFFFFF", size="0")
    for index, width in enumerate((4680, 4680)):
        set_cell_width(wrapper.cell(0, index), width)
        set_cell_margins(wrapper.cell(0, index), top=80, bottom=80, start=40, end=40)

    earnings_cell = wrapper.cell(0, 0)
    deductions_cell = wrapper.cell(0, 1)
    add_repeating_table_in_cell(
        earnings_cell,
        "Earnings",
        "PayslipEarnings",
        [("Description", "Description", 3300, "left"), ("Amount", "Amount", 1380, "right")],
    )
    add_repeating_table_in_cell(
        deductions_cell,
        "Deductions",
        "PayslipDeductions",
        [("Description", "Description", 3300, "left"), ("Amount", "Amount", 1380, "right")],
    )


def add_repeating_table_in_cell(cell, title, group_name, columns):
    cell.text = ""
    p = cell.paragraphs[0]
    add_static_text(p, title, size=9.5, color=NAVY, bold=True)
    table = cell.add_table(rows=2, cols=len(columns))
    table.autofit = False
    set_table_borders(table)
    for index, (header, field, width, align) in enumerate(columns):
        header_cell = table.cell(0, index)
        set_cell_width(header_cell, width)
        set_cell_fill(header_cell, HEADER_FILL)
        set_cell_margins(header_cell, top=70, bottom=70, start=80, end=80)
        p = clear_cell(header_cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_static_text(p, header, size=7.5, color=NAVY, bold=True)

        value_cell = table.cell(1, index)
        set_cell_width(value_cell, width)
        set_cell_margins(value_cell, top=70, bottom=70, start=80, end=80)
        p = clear_cell(value_cell)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == "right" else WD_ALIGN_PARAGRAPH.LEFT
        if index == 0:
            add_merge_field(p, f"TableStart:{group_name}", display="", size=1, hidden=True)
        add_merge_field(p, field, size=7.7)
        if index == len(columns) - 1:
            add_merge_field(p, f"TableEnd:{group_name}", display="", size=1, hidden=True)


def add_totals_and_notes(doc, config):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table)
    set_table_borders(table, color="FFFFFF", size="0")
    set_cell_width(table.cell(0, 0), 5800)
    set_cell_width(table.cell(0, 1), 3560)
    for cell in table.row_cells(0):
        set_cell_margins(cell, top=80, bottom=80, start=80, end=80)

    notes = table.cell(0, 0)
    p = clear_cell(notes)
    add_static_text(p, "Notes", size=9.5, color=NAVY, bold=True)
    p = notes.add_paragraph()
    add_merge_field(p, "Notes", size=8)

    if config.get("payment"):
        p = notes.add_paragraph()
        style_paragraph(p, before=4, after=1)
        add_static_text(p, "Payment / Bank Details", size=9, color=NAVY, bold=True)
        p = notes.add_paragraph()
        add_static_text(
            p,
            "Editable placeholder: add bank name, account name, account number, branch, SWIFT, or payment instructions here.",
            size=8,
            color=MUTED,
            italic=True,
        )

    p = notes.add_paragraph()
    style_paragraph(p, before=4, after=1)
    add_static_text(p, "Terms", size=9, color=NAVY, bold=True)
    p = notes.add_paragraph()
    add_static_text(p, config["terms"], size=8, color=MUTED)

    totals = table.cell(0, 1).add_table(rows=4, cols=2)
    set_table_width(totals, width_twips=3300)
    set_table_borders(totals)
    for row_index, (label, field) in enumerate(
        [("Sub Total", "SubTotal"), ("Tax Total", "TaxTotal"), ("Discount / Deduction", "DiscountTotal"), ("Grand / Net Total", "GrandTotal")]
    ):
        row = totals.rows[row_index]
        for index, width in enumerate((1650, 1650)):
            cell = row.cells[index]
            set_cell_width(cell, width)
            set_cell_margins(cell, top=70, bottom=70, start=100, end=100)
            if row_index == 3:
                set_cell_fill(cell, "0B2545")
        p = clear_cell(row.cells[0])
        add_static_text(p, label, size=8.2, color=WHITE if row_index == 3 else MUTED, bold=True)
        p = clear_cell(row.cells[1])
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_merge_field(p, field, size=8.2, color=WHITE if row_index == 3 else NAVY, bold=True)
        if row_index == 3:
            add_static_text(p, " ", size=8.2, color=WHITE)
            add_merge_field(p, "CurrencyCode", size=8.2, color=WHITE, bold=True)


def add_signatures(doc, config):
    table = doc.add_table(rows=2, cols=2)
    set_table_width(table)
    set_table_borders(table, color="FFFFFF", size="0")
    for row in table.rows:
        for cell in row.cells:
            set_cell_width(cell, 4680)
            set_cell_margins(cell, top=80, bottom=40, start=120, end=120)

    for index, label in enumerate(config["signatures"]):
        p = clear_cell(table.cell(0, index))
        style_paragraph(p, before=8, after=0)
        add_static_text(p, "_" * 34, size=8, color=MUTED)
        p = clear_cell(table.cell(1, index))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_static_text(p, label, size=8, color=MUTED, bold=True)


def add_footer(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_static_text(
        p,
        "Generated by NexusFlow ERP | Template fields are populated from Company Settings and print document data.",
        size=7,
        color=MUTED,
    )


def build_template(config):
    doc = Document()
    configure_document(doc)
    add_header(doc)
    add_title_block(doc, config)
    add_address_block(doc, config)
    if config.get("payroll"):
        add_payslip_tables(doc)
    else:
        add_repeating_table(doc, config["table_title"], config["table_name"], config["columns"])
    add_totals_and_notes(doc, config)
    add_signatures(doc, config)
    add_footer(doc)
    return doc


def write_readme():
    lines = [
        "# NexusFlow Default Document Templates",
        "",
        "This folder contains editable A4 portrait `.docx` templates for the NexusFlow reusable print preview engine.",
        "",
        "## How to use",
        "",
        "1. Open **Company Settings**.",
        "2. Go to the **Document Templates** tab.",
        "3. Choose the document type and tax profile.",
        "4. Upload the matching `.docx` file from this folder.",
        "5. Mark the template as default when it should be used for generated PDFs.",
        "",
        "## Common merge fields",
        "",
        "- Company fields: `CompanyName`, `CompanyTaxRegistrationNumber`, `CompanyBusinessRegistrationNumber`, `CompanyAddress`, `CompanyEmail`, `CompanyPhone`, `Image:CompanyLogo`",
        "- Header fields: `DocumentNumber`, `DocumentDate`, `CustomerOrSupplierName`, `BillingAddress`, `ShippingAddress`, `Notes`, `CurrencyCode`",
        "- Totals: `SubTotal`, `TaxTotal`, `DiscountTotal`, `GrandTotal`",
        "",
        "## Repeating groups by document type",
        "",
    ]

    for config in DOCUMENTS:
        if config.get("payroll"):
            lines.append(f"- `{config['type']}`: `PayslipEarnings` and `PayslipDeductions` with `Description`, `Amount`.")
        else:
            fields = ", ".join(f"`{field}`" for _, field, _, _ in config["columns"])
            lines.append(f"- `{config['type']}`: `{config['table_name']}` with {fields}.")

    lines.extend(
        [
            "",
            "Payment and bank-detail text is intentionally editable static placeholder content because NexusFlow does not yet expose dynamic bank-detail fields in `PrintDocumentDto`.",
            "Paysheet is not generated here because the current application enum contains `Payslip` but not a separate `Paysheet` document type.",
            "",
        ]
    )

    (REVIEW_DIR / "README.md").write_text("\n".join(lines), encoding="utf-8")


def main():
    REVIEW_DIR.mkdir(parents=True, exist_ok=True)
    SEED_DIR.mkdir(parents=True, exist_ok=True)

    for config in DOCUMENTS:
        doc = build_template(config)
        file_name = f"{config['type']}_Default.docx"
        review_path = REVIEW_DIR / file_name
        seed_path = SEED_DIR / file_name
        doc.save(review_path)
        shutil.copyfile(review_path, seed_path)

    write_readme()
    print(f"Generated {len(DOCUMENTS)} templates in {REVIEW_DIR}")
    print(f"Copied templates to {SEED_DIR}")


if __name__ == "__main__":
    main()
